import os
import shutil
import sys

# Ensure backend folder is in path so we can import parsers
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

from parsers import parse_document
from parsers.url_parser import parse_url

# Test constants and paths
TEMP_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "temp", "test_files"))
TXT_PATH = os.path.join(TEMP_DIR, "sample.txt")
DOCX_PATH = os.path.join(TEMP_DIR, "sample.docx")
PDF_PATH = os.path.join(TEMP_DIR, "sample.pdf")

def setup_test_environment():
    """Create a temporary directory for test files if it doesn't exist."""
    if not os.path.exists(TEMP_DIR):
        os.makedirs(TEMP_DIR)

def generate_sample_txt():
    """Generates a simple sample text file."""
    content = "   This is a sample requirements text file.\nREQ-001: The system shall run offline.\nREQ-002: Latency must be low.   "
    with open(TXT_PATH, "w", encoding="utf-8") as f:
        f.write(content)
    return content.strip()

def generate_sample_docx():
    """Generates a sample DOCX document containing paragraphs and a table using python-docx."""
    from docx import Document
    doc = Document()
    
    # Add paragraphs
    doc.add_paragraph("Requirements Document v1.0")
    doc.add_paragraph("This paragraph contains system specification details.")
    
    # Add a table
    table = doc.add_table(rows=3, cols=3)
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "ID"
    hdr_cells[1].text = "Requirement"
    hdr_cells[2].text = "Type"
    
    # Row 1
    row_1 = table.rows[1].cells
    row_1[0].text = "REQ-101"
    row_1[1].text = "File parsing must support PDF format."
    row_1[2].text = "Mechanism"
    
    # Row 2
    row_2 = table.rows[2].cells
    row_2[0].text = "REQ-102"
    row_2[1].text = "File parsing must support DOCX format."
    row_2[2].text = "Mechanism"
    
    doc.save(DOCX_PATH)

def generate_sample_pdf():
    """Generates a sample PDF document containing paragraphs and table-like text using reportlab."""
    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import letter
    
    c = canvas.Canvas(PDF_PATH, pagesize=letter)
    
    # Add page 1 content
    c.setFont("Helvetica-Bold", 16)
    c.drawString(72, 750, "Requirements Document - PDF")
    
    c.setFont("Helvetica", 10)
    c.drawString(72, 720, "This is page 1 of the requirements document.")
    c.drawString(72, 700, "Below is a simulated specifications table:")
    
    # Draw simple table rows
    c.drawString(72, 670, "ID | Requirement | Category")
    c.drawString(72, 650, "REQ-201 | The system shall parse PDF tables | Mechanism")
    c.drawString(72, 630, "REQ-202 | Parsing must complete within 2 seconds | Performance")
    
    c.showPage()
    
    # Add page 2 content
    c.setFont("Helvetica", 10)
    c.drawString(72, 750, "This is page 2 of the requirements document.")
    c.drawString(72, 730, "Ensure that page breaks do not crash the parser.")
    
    c.save()

def cleanup_test_environment():
    """Removes all generated files and temporary test directories."""
    if os.path.exists(TEMP_DIR):
        shutil.rmtree(TEMP_DIR)

def run_tests():
    setup_test_environment()
    
    results = {
        "txt": "FAILED",
        "docx": "FAILED",
        "pdf": "FAILED",
        "url": "FAILED"
    }
    
    try:
        # 1. Test TXT Parser
        print("--- Testing TXT Parser ---")
        expected_txt = generate_sample_txt()
        parsed_txt = parse_document(TXT_PATH, "txt")
        print("Parsed Text Content (First 500 chars):")
        print("-" * 40)
        print(parsed_txt[:500])
        print("-" * 40)
        
        if parsed_txt == expected_txt:
            results["txt"] = "PASSED"
            print("TXT Parser Status: PASS")
        else:
            print(f"TXT Parser Status: FAIL (content mismatch)\nExpected:\n{expected_txt}\nGot:\n{parsed_txt}")
        print()

        # 2. Test DOCX Parser
        print("--- Testing DOCX Parser ---")
        generate_sample_docx()
        parsed_docx = parse_document(DOCX_PATH, "docx")
        print("Parsed DOCX Content (First 500 chars):")
        print("-" * 40)
        print(parsed_docx[:500])
        print("-" * 40)
        
        # Verify text and table presence
        if "Requirements Document v1.0" in parsed_docx and "ID | Requirement | Type" in parsed_docx and "REQ-101 | File parsing must support PDF format. | Mechanism" in parsed_docx:
            results["docx"] = "PASSED"
            print("DOCX Parser Status: PASS")
        else:
            print("DOCX Parser Status: FAIL (missing expected paragraphs or table rows)")
        print()

        # 3. Test PDF Parser
        print("--- Testing PDF Parser ---")
        generate_sample_pdf()
        parsed_pdf = parse_document(PDF_PATH, "pdf")
        print("Parsed PDF Content (First 500 chars):")
        print("-" * 40)
        print(parsed_pdf[:500])
        print("-" * 40)
        
        # Verify text and table presence
        if "Requirements Document - PDF" in parsed_pdf and "ID | Requirement | Category" in parsed_pdf and "REQ-201 | The system shall parse PDF tables | Mechanism" in parsed_pdf:
            results["pdf"] = "PASSED"
            print("PDF Parser Status: PASS")
        else:
            print("PDF Parser Status: FAIL (missing expected text or simulated table rows)")
        print()

        # 4. Test URL Parser
        print("--- Testing URL Parser ---")
        test_url = "https://example.com"
        print(f"Fetching and parsing {test_url}...")
        parsed_url = parse_url(test_url)
        print("Parsed URL Content (First 500 chars):")
        print("-" * 40)
        print(parsed_url[:500])
        print("-" * 40)
        
        # Verify we extracted text from example.com
        if "Example Domain" in parsed_url:
            results["url"] = "PASSED"
            print("URL Parser Status: PASS")
        else:
            print("URL Parser Status: FAIL (could not find 'Example Domain' in output)")
        print()

    except Exception as e:
        print(f"\nAn error occurred during testing: {str(e)}")
        import traceback
        traceback.print_exc()
        
    finally:
        cleanup_test_environment()
        
    print("=" * 40)
    print("PARSER TEST SUMMARY:")
    print("=" * 40)
    all_passed = True
    for parser, status in results.items():
        print(f"{parser.upper()} Parser: {status}")
        if status != "PASSED":
            all_passed = False
    print("=" * 40)
    
    if all_passed:
        print("ALL TESTS PASSED SUCCESSFULLY!")
        sys.exit(0)
    else:
        print("SOME TESTS FAILED.")
        sys.exit(1)

if __name__ == "__main__":
    run_tests()
